import streamlit as st
from docx import Document
import tempfile
import os
import json
import fitz
import msoffcrypto
from io import BytesIO
from langdetect import detect
from datetime import datetime
import requests

st.set_page_config(page_title="Dịch Trung - Việt (Từ điển + Google Requests)", layout="centered")

def decrypt_office_file(uploaded_file):
    decrypted = BytesIO()
    try:
        office_file = msoffcrypto.OfficeFile(uploaded_file)
        office_file.load_key(password=None)
        office_file.decrypt(decrypted)
        decrypted.seek(0)
        return decrypted
    except Exception:
        return None

def read_pdf(file_bytes):
    try:
        doc = fitz.open(stream=file_bytes.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    except Exception:
        return ""

def read_txt(file_bytes):
    try:
        return file_bytes.read().decode("utf-8")
    except Exception:
        return ""

def detect_language(text):
    try:
        return detect(text)
    except Exception:
        return "unknown"

def load_dictionary():
    try:
        with open("dictionary_full.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        st.error("Không thể đọc file từ điển.")
        return {}

def translate_by_dictionary(sentence, dictionary):
    result = []
    translation_result = []
    i = 0
    while i < len(sentence):
        found = False
        for j in range(min(5, len(sentence) - i), 0, -1):
            segment = sentence[i:i+j]
            if segment in dictionary:
                result.append(f"{segment}: {dictionary[segment]}")
                translation_result.append(dictionary[segment])
                i += j
                found = True
                break
        if not found:
            char = sentence[i]
            result.append(f"{char}: Không rõ nghĩa")
            translation_result.append(char)
            i += 1
    return "\n".join(result), "→ Dịch toàn câu (từ điển): " + " ".join(translation_result)

def translate_by_google(text):
    try:
        url = "https://translate.googleapis.com/translate_a/single"
        params = {
            "client": "gtx",
            "sl": "zh-CN",
            "tl": "vi",
            "dt": "t",
            "q": text
        }
        r = requests.get(url, params=params)
        r.raise_for_status()
        result = r.json()
        return "".join([seg[0] for seg in result[0]])
    except Exception:
        return "[Google dịch thất bại]"


def clean_text_for_docx(text):
    if not isinstance(text, str):
        text = str(text)
    # Xoá ký tự điều khiển không hợp lệ trong XML
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)


def create_word_doc(sentences, dictionary, method):
    doc = Document()
    for zh in sentences:
        if zh.strip() == "":
            continue
        safe_zh = zh.encode('utf-8', 'ignore').decode('utf-8')
        doc.add_paragraph(clean_text_for_docx(safe_zh))
        word_translation, sentence_meaning = translate_by_dictionary(zh, dictionary)
        doc.add_paragraph(clean_text_for_docx(word_translation))
        if method == "Từ điển":
            doc.add_paragraph(clean_text_for_docx(sentence_meaning))
        elif method == "Google Translate":
            google_result = translate_by_google(zh)
            doc.add_paragraph(clean_text_for_docx("→ Google dịch: " + google_result))
    return doc

st.title("📘 Dịch Trung - Việt Song Ngữ")

method = st.radio("Phương thức dịch câu:", ["Từ điển", "Google Translate"])
uploaded_file = st.file_uploader("📎 Tải lên file văn bản (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"])

dictionary = load_dictionary()

if uploaded_file:
    with st.spinner("🔍 Đang xử lý..."):
        ext = uploaded_file.name.split(".")[-1].lower()
        content = ""
        if ext == "txt":
            content = read_txt(uploaded_file)
        elif ext == "pdf":
            content = read_pdf(uploaded_file)
        elif ext == "docx":
            decrypted = decrypt_office_file(uploaded_file)
            if decrypted:
                docx_obj = Document(decrypted)
                content = "\n".join([p.text for p in docx_obj.paragraphs])
            else:
                st.error("❌ Không thể mở khóa file Word.")
                st.stop()

        if not content.strip():
            st.warning("⚠️ File không chứa nội dung hợp lệ.")
            st.stop()

        lang = detect_language(content)
        st.info(f"🌐 Ngôn ngữ phát hiện: {lang}")

        sentences = [line.strip() for line in content.strip().splitlines() if line.strip()]
        doc = create_word_doc(sentences, dictionary, method)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            st.success("✅ Hoàn tất dịch!")
            st.download_button(
                label="📥 Tải xuống bản Word song ngữ",
                data=open(tmp.name, "rb").read(),
                file_name="trung_viet_final.docx"
            )
            os.unlink(tmp.name)
